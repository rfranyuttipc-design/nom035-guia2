"""
generar_reporte_g2.py
=====================
Módulo de reportes avanzados NOM-035-STPS-2018 — Guía II
Genera:
  1. Excel mejorado con análisis completo (mapa de calor, dominios, violencia)
  2. Informe Word profesional (Tabla 5, programa de acción, informe STPS)
"""

import os, io, textwrap
from datetime import datetime

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from openpyxl import load_workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Escalas de calificación NOM-035 Guía II (Tabla 2) ─────────────────────────
ESCALA_D = {"Siempre":4,"Casi siempre":3,"Algunas veces":2,"Casi nunca":1,"Nunca":0}
ESCALA_I = {"Siempre":0,"Casi siempre":1,"Algunas veces":2,"Casi nunca":3,"Nunca":4}
ITEMS_INV_G2 = {18,19,20,21,22,23,24,25,26,27,28,29,30,31,32,33}

# ── Paleta ────────────────────────────────────────────────────────────────────
VERDE       = "4B694E"
VERDE_CLARO = "D6E4D8"
VERDE_MID   = "6A9370"
ROJO        = "A20000"
ROJO_CLARO  = "FDDEDE"
NARANJA     = "E07820"
AMARILLO_BG = "FFF8DC"
AZUL        = "69A2D8"
AZUL_CLARO  = "EBF3FB"
AMBAR       = "C8A600"
GRIS        = "F4F4F4"
BLANCO      = "FFFFFF"

C_NULO     = "#4B694E"
C_BAJO     = "#69A2D8"
C_MEDIO    = "#C8A600"
C_ALTO     = "#E07820"
C_MUY_ALTO = "#A20000"

NIVEL_COLORES = {
    "NULO":     (C_NULO,     "#D6E4D8"),
    "BAJO":     (C_BAJO,     "#EBF3FB"),
    "MEDIO":    (C_MEDIO,    "#FFF8DC"),
    "ALTO":     (C_ALTO,     "#FDEBD0"),
    "MUY ALTO": (C_MUY_ALTO, "#FDDEDE"),
}

NOMBRES_DOM = {
    0: "Condiciones del ambiente",
    1: "Carga de trabajo",
    2: "Falta de control",
    3: "Jornada de trabajo",
    4: "Interferencia trabajo-familia",
    5: "Liderazgo",
    6: "Relaciones en el trabajo",
    7: "Violencia laboral",
}

TABLA5 = [
    (0,  20,  "NULO"),
    (21, 45,  "BAJO"),
    (46, 70,  "MEDIO"),
    (71, 90,  "ALTO"),
    (91, 999, "MUY ALTO"),
]

# ── Helpers Excel ─────────────────────────────────────────────────────────────
def _fill(h): return PatternFill("solid", fgColor=h)
def _font(bold=False, size=10, color="1A1A1A", name="Calibri"):
    return Font(bold=bold, size=size, color=color, name=name)
def _align(h="left", v="center", wrap=False):
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)
def _side(c="D0D0D0"): return Side(border_style="thin", color=c)
def _border(c="D0D0D0"):
    s = _side(c)
    return Border(left=s, right=s, top=s, bottom=s)

def _auto_width(ws, mn=8, mx=40):
    for col in ws.columns:
        ml = max((len(str(c.value or "")) for c in col), default=0)
        ws.column_dimensions[get_column_letter(col[0].column)].width = min(max(ml+2, mn), mx)

def _buf(fig):
    b = io.BytesIO()
    fig.savefig(b, format="png", dpi=180, bbox_inches="tight", facecolor="white")
    b.seek(0); plt.close(fig); return b

# ── Calcular nivel desde puntaje ──────────────────────────────────────────────
def _nivel(p):
    for pmin, pmax, nv in TABLA5:
        if pmin <= p <= pmax:
            return nv
    return "MUY ALTO"

# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICAS
# ══════════════════════════════════════════════════════════════════════════════
def _graf_distribucion_riesgo(df):
    """Gráfica de pastel — distribución por nivel de riesgo global."""
    if "Nivel Riesgo" not in df.columns:
        return None
    conteo = df["Nivel Riesgo"].value_counts()
    orden  = ["NULO","BAJO","MEDIO","ALTO","MUY ALTO"]
    labels = [n for n in orden if n in conteo.index]
    vals   = [conteo[n] for n in labels]
    cols   = [NIVEL_COLORES.get(n, ("#888888","#eeeeee"))[0] for n in labels]
    fig, ax = plt.subplots(figsize=(5.5, 4.5))
    wedges, texts, autotexts = ax.pie(
        vals, labels=labels, autopct="%1.1f%%", colors=cols,
        startangle=90, wedgeprops=dict(edgecolor="white", linewidth=1.5))
    for t in texts:     t.set_fontsize(8)
    for a in autotexts: a.set_fontsize(8); a.set_color("white"); a.set_fontweight("bold")
    ax.set_title("Distribución por Nivel de Riesgo (Tabla 5 NOM-035)",
                 fontsize=10, fontweight="bold", pad=10)
    fig.tight_layout()
    return _buf(fig)

def _graf_mapa_calor_dominios(df):
    """Mapa de calor por dominio — promedio de puntaje por dominio."""
    dom_cols = {
        0: "Dom0 Ambiente", 1: "Dom1 Carga", 2: "Dom2 Control",
        3: "Dom3 Jornada",  4: "Dom4 Interferencia", 5: "Dom5 Liderazgo",
        6: "Dom6 Relaciones", 7: "Dom7 Violencia"
    }
    promedios = {}
    for dom_id, col in dom_cols.items():
        if col in df.columns:
            promedios[NOMBRES_DOM[dom_id]] = df[col].mean()

    if not promedios:
        return None

    nombres = list(promedios.keys())
    valores = list(promedios.values())
    # Color por nivel (aproximado)
    maximos = [16, 48, 24, 12, 20, 36, 16, 12]
    pcts    = [v/m*100 if m > 0 else 0 for v, m in zip(valores, maximos)]
    cols    = []
    for p in pcts:
        if p < 25:   cols.append(C_NULO)
        elif p < 50: cols.append(C_BAJO)
        elif p < 70: cols.append(C_MEDIO)
        elif p < 85: cols.append(C_ALTO)
        else:        cols.append(C_MUY_ALTO)

    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.barh(nombres[::-1], valores[::-1], color=cols[::-1],
                   edgecolor="white", height=0.6)
    for bar, val, pct in zip(bars, valores[::-1], pcts[::-1]):
        ax.text(bar.get_width() + 0.2, bar.get_y() + bar.get_height()/2,
                f"{val:.1f} pts ({pct:.0f}%)", va="center", fontsize=8)
    ax.set_xlabel("Puntaje promedio", fontsize=9)
    ax.set_title("Mapa de Dominios — Puntaje Promedio por Dominio",
                 fontsize=10, fontweight="bold", pad=8)
    ax.spines[["top","right"]].set_visible(False)

    # Leyenda semáforo
    parches = [
        mpatches.Patch(color=C_NULO,     label="NULO (<25%)"),
        mpatches.Patch(color=C_BAJO,     label="BAJO (25-49%)"),
        mpatches.Patch(color=C_MEDIO,    label="MEDIO (50-69%)"),
        mpatches.Patch(color=C_ALTO,     label="ALTO (70-84%)"),
        mpatches.Patch(color=C_MUY_ALTO, label="MUY ALTO (≥85%)"),
    ]
    ax.legend(handles=parches, fontsize=7, loc="lower right")
    fig.tight_layout()
    return _buf(fig)

def _graf_radar_dominios(df):
    """Gráfica de radar por dominio."""
    dom_cols = {
        "Ambiente":      "Dom0 Ambiente",
        "Carga":         "Dom1 Carga",
        "Control":       "Dom2 Control",
        "Jornada":       "Dom3 Jornada",
        "Interf.Familia":"Dom4 Interferencia",
        "Liderazgo":     "Dom5 Liderazgo",
        "Relaciones":    "Dom6 Relaciones",
        "Violencia":     "Dom7 Violencia",
    }
    maximos = [16, 48, 24, 12, 20, 36, 16, 12]
    nombres  = list(dom_cols.keys())
    promedios = []
    for col, maximo in zip(dom_cols.values(), maximos):
        if col in df.columns:
            promedios.append(min(df[col].mean() / maximo, 1.0))
        else:
            promedios.append(0)

    N = len(nombres)
    angulos = [n / float(N) * 2 * 3.14159 for n in range(N)]
    angulos += angulos[:1]
    promedios += promedios[:1]

    fig, ax = plt.subplots(figsize=(5.5, 5.5), subplot_kw=dict(polar=True))
    ax.set_theta_offset(3.14159 / 2)
    ax.set_theta_direction(-1)
    ax.set_xticks(angulos[:-1])
    ax.set_xticklabels(nombres, fontsize=8)
    ax.set_ylim(0, 1)
    ax.set_yticks([0.25, 0.5, 0.75, 1.0])
    ax.set_yticklabels(["25%","50%","75%","100%"], fontsize=7)
    ax.plot(angulos, promedios, "o-", linewidth=2, color="#4B694E")
    ax.fill(angulos, promedios, alpha=0.25, color="#4B694E")
    ax.set_title("Radar de Dominios\n(% del puntaje máximo)",
                 fontsize=10, fontweight="bold", pad=15)
    fig.tight_layout()
    return _buf(fig)

def _graf_violencia(df):
    """Gráfica violencia laboral — ítems 33-40 NOM-035 Guía II (Tabla 3)."""
    escala_d = {"Siempre":4,"Casi siempre":3,"Algunas veces":2,"Casi nunca":1,"Nunca":0}
    escala_i = {"Siempre":0,"Casi siempre":1,"Algunas veces":2,"Casi nunca":3,"Nunca":4}
    cols_v  = ["P33","P34","P35","P36","P37","P38","P39","P40"]
    present = [c for c in cols_v if c in df.columns]
    if not present:
        return None
    labels = {
        "P33":"Sin expresión libre (inv.)",
        "P34":"Críticas constantes",
        "P35":"Burlas, humillaciones",
        "P36":"Me excluyen",
        "P37":"Me hacen quedar mal",
        "P38":"Ignoran mis logros",
        "P39":"Bloquean ascensos",
        "P40":"Presencié violencia",
    }
    escala_d = {"Siempre":4,"Casi siempre":3,"Algunas veces":2,"Casi nunca":1,"Nunca":0}
    escala_i = {"Siempre":0,"Casi siempre":1,"Algunas veces":2,"Casi nunca":3,"Nunca":4}
    etiquetas, promedios = [], []
    for col in present:
        esc = escala_i if col == "P33" else escala_d
        vals = df[col].map(esc).fillna(0)
        promedios.append(vals.mean())
        etiquetas.append(labels.get(col, col))
    fig, ax = plt.subplots(figsize=(7, 4.5))
    bars = ax.barh(etiquetas[::-1], promedios[::-1],
                   color=[C_MUY_ALTO if p > 2 else C_ALTO if p > 1 else
                          C_MEDIO if p > 0.5 else C_NULO for p in promedios[::-1]],
                   edgecolor="white", height=0.55)
    for bar, val in zip(bars, promedios[::-1]):
        ax.text(bar.get_width()+0.05, bar.get_y()+bar.get_height()/2,
                f"{val:.2f}/4", va="center", fontsize=8)
    ax.set_xlim(0, 4.8)
    ax.set_xlabel("Puntaje promedio de riesgo (0=Sin riesgo, 4=Máximo)", fontsize=9)
    ax.set_title("Análisis de Violencia Laboral — ítems 33-40 (NOM-035 Guía II)",
                 fontsize=10, fontweight="bold", pad=8)
    ax.spines[["top","right"]].set_visible(False)
    ax.axvline(x=1, color="orange", linestyle="--", alpha=0.5, linewidth=1, label="Umbral atención")
    ax.axvline(x=2, color="red",    linestyle="--", alpha=0.4, linewidth=1, label="Umbral urgente")
    ax.legend(fontsize=8)
    fig.tight_layout()
    return _buf(fig)


def _graf_por_area(df):
    """Riesgo promedio por área."""
    if "Área" not in df.columns or "Puntaje Total" not in df.columns:
        return None
    por_area = df.groupby("Área")["Puntaje Total"].mean().sort_values(ascending=False)
    if len(por_area) < 1:
        return None
    fig, ax = plt.subplots(figsize=(7, max(3, len(por_area)*0.6)))
    cols = [C_MUY_ALTO if v > 90 else C_ALTO if v > 70 else
            C_MEDIO if v > 45 else C_BAJO if v > 20 else C_NULO
            for v in por_area.values]
    bars = ax.barh(por_area.index[::-1], por_area.values[::-1],
                   color=cols[::-1], edgecolor="white", height=0.6)
    for bar, val in zip(bars, por_area.values[::-1]):
        nivel = _nivel(int(val))
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                f"{val:.1f} pts — {nivel}", va="center", fontsize=8)
    ax.set_xlabel("Puntaje promedio", fontsize=9)
    ax.set_title("Puntaje Promedio por Área (Foco de Atención)",
                 fontsize=10, fontweight="bold", pad=8)
    ax.spines[["top","right"]].set_visible(False)
    # Líneas de referencia
    for x, lbl, clr in [(20,"NULO","#4B694E"),(45,"BAJO","#69A2D8"),
                         (70,"MEDIO","#C8A600"),(90,"ALTO","#E07820")]:
        ax.axvline(x=x, color=clr, linestyle="--", alpha=0.4, linewidth=0.8)
    fig.tight_layout()
    return _buf(fig)

def _graf_carga_trabajo(df):
    """Análisis detallado ítems carga de trabajo."""
    escala_d = {"Siempre":4,"Casi siempre":3,"Algunas veces":2,"Casi nunca":1,"Nunca":0}
    items_carga = {
        "P04":"Tiempo extra",
        "P05":"Trabajar sin parar",
        "P06":"Ritmo acelerado",
        "P07":"Alta concentración",
        "P08":"Memoriza mucha info",
        "P09":"Varios asuntos a la vez",
        "P10":"Resp. cosas valiosas",
        "P11":"Resp. área completa",
        "P12":"Órdenes contradictorias",
        "P13":"Cosas innecesarias",
    }
    escala = {"Siempre":4,"Casi siempre":3,"Algunas veces":2,"Casi nunca":1,"Nunca":0}
    promedios = {}
    for col, etiq in items_carga.items():
        if col in df.columns:
            promedios[etiq] = df[col].map(escala).fillna(0).mean()

    if not promedios:
        return None

    fig, ax = plt.subplots(figsize=(7, 5))
    etiquetas = list(promedios.keys())
    valores   = list(promedios.values())
    cols = [C_MUY_ALTO if v > 3 else C_ALTO if v > 2 else
            C_MEDIO if v > 1 else C_BAJO for v in valores]
    bars = ax.barh(etiquetas[::-1], valores[::-1], color=cols[::-1],
                   edgecolor="white", height=0.6)
    for bar, val in zip(bars, valores[::-1]):
        ax.text(bar.get_width() + 0.05, bar.get_y() + bar.get_height()/2,
                f"{val:.2f}", va="center", fontsize=8)
    ax.set_xlim(0, 4.5)
    ax.set_xlabel("Promedio (0=Nunca, 4=Siempre)", fontsize=9)
    ax.set_title("Análisis Detallado — Carga de Trabajo (ítems 5-16)",
                 fontsize=10, fontweight="bold", pad=8)
    ax.spines[["top","right"]].set_visible(False)
    fig.tight_layout()
    return _buf(fig)

# ══════════════════════════════════════════════════════════════════════════════
# ANALÍTICAS AVANZADAS
# ══════════════════════════════════════════════════════════════════════════════
def _calcular_analiticas_g2(df: pd.DataFrame) -> dict:
    """Calcula todos los indicadores avanzados de la Guía II."""
    total = len(df)
    if total == 0:
        return {}

    escala_d = {"Siempre":4,"Casi siempre":3,"Algunas veces":2,"Casi nunca":1,"Nunca":0}
    escala_i = {"Siempre":0,"Casi siempre":1,"Algunas veces":2,"Casi nunca":3,"Nunca":4}
    escala   = escala_d

    # 1. Distribución por nivel de riesgo
    dist_riesgo = {}
    if "Nivel Riesgo" in df.columns:
        for nv in ["NULO","BAJO","MEDIO","ALTO","MUY ALTO"]:
            n = (df["Nivel Riesgo"] == nv).sum()
            dist_riesgo[nv] = {"n": int(n), "pct": round(n/total*100, 1)}

    # Casos que requieren intervención
    casos_alto    = dist_riesgo.get("ALTO",     {"n":0})["n"]
    casos_muy_alto= dist_riesgo.get("MUY ALTO", {"n":0})["n"]
    casos_criticos= casos_alto + casos_muy_alto

    # 2. Puntaje promedio por dominio
    dom_cols = {
        0:"Dom0 Ambiente", 1:"Dom1 Carga", 2:"Dom2 Control",
        3:"Dom3 Jornada",  4:"Dom4 Interferencia", 5:"Dom5 Liderazgo",
        6:"Dom6 Relaciones", 7:"Dom7 Violencia"
    }
    maximos_dom = {0:16, 1:48, 2:24, 3:12, 4:20, 5:36, 6:16, 7:12}
    promedios_dom = {}
    niveles_dom   = {}
    for dom_id, col in dom_cols.items():
        if col in df.columns:
            prom = df[col].mean()
            pct  = prom / maximos_dom[dom_id] * 100
            promedios_dom[dom_id] = round(prom, 1)
            if pct < 25:   nv = "NULO"
            elif pct < 50: nv = "BAJO"
            elif pct < 70: nv = "MEDIO"
            elif pct < 85: nv = "ALTO"
            else:          nv = "MUY ALTO"
            niveles_dom[dom_id] = nv

    # Dominio más crítico
    if promedios_dom:
        dom_critico_id  = max(promedios_dom,
                              key=lambda d: promedios_dom[d]/maximos_dom[d])
        dom_critico_nom = NOMBRES_DOM[dom_critico_id]
        dom_critico_niv = niveles_dom.get(dom_critico_id, "")
    else:
        dom_critico_id  = 0
        dom_critico_nom = "N/D"
        dom_critico_niv = ""

    # 3. Violencia laboral (ítems 44-46)
    cols_viol = ["P34","P35","P36","P37","P38","P39","P40"]
    n_viol    = 0
    viol_prom = {}
    for col in cols_viol:
        if col in df.columns:
            esc  = escala_i if col == "P33" else escala_d
            vals = df[col].map(esc).fillna(0)
            viol_prom[col] = round(vals.mean(), 2)
            n_viol += (vals > 0).sum()
    alertas_violencia = int((df.get("Alerta Violencia","No") == "SÍ — URGENTE").sum()) \
                        if "Alerta Violencia" in df.columns else 0

    # 4. Segmentación por área
    seg_area = {}
    if "Área" in df.columns and "Nivel Riesgo" in df.columns:
        for area, grp in df.groupby("Área"):
            n_area  = len(grp)
            n_alto  = ((grp["Nivel Riesgo"] == "ALTO") | (grp["Nivel Riesgo"] == "MUY ALTO")).sum()
            prom_pt = grp["Puntaje Total"].mean() if "Puntaje Total" in grp.columns else 0
            seg_area[area] = {
                "total":   int(n_area),
                "criticos": int(n_alto),
                "pct":     round(n_alto/n_area*100, 1),
                "puntaje_prom": round(prom_pt, 1),
                "nivel_prom":  _nivel(int(prom_pt)),
            }
    foco_rojo = max(seg_area.items(), key=lambda x: x[1]["pct"])[0] \
                if seg_area else "N/D"

    # 5. Correlaciones
    correlaciones = []
    # Jornada extensa vs Violencia
    if "Dom3 Jornada" in df.columns and "Dom7 Violencia" in df.columns:
        jornada_alta  = df["Dom3 Jornada"] > df["Dom3 Jornada"].median()
        viol_alta     = df["Dom7 Violencia"] > 0
        overlap = (jornada_alta & viol_alta).sum()
        if overlap > 0:
            correlaciones.append(
                f"{overlap} trabajador(es) con jornada extensa también reportan violencia laboral."
            )
    # Liderazgo vs Interferencia trabajo-familia
    if "Dom5 Liderazgo" in df.columns and "Dom4 Interferencia" in df.columns:
        liderazgo_bajo = df["Dom5 Liderazgo"] < df["Dom5 Liderazgo"].quantile(0.25)
        interf_alta    = df["Dom4 Interferencia"] > df["Dom4 Interferencia"].median()
        overlap2 = (liderazgo_bajo & interf_alta).sum()
        if overlap2 > 0:
            correlaciones.append(
                f"{overlap2} trabajador(es) con bajo apoyo de liderazgo reportan alta "
                f"interferencia trabajo-familia — posible relación directa."
            )

    # 6. Párrafo ejecutivo
    nivel_predominante = max(dist_riesgo.items(),
                             key=lambda x: x[1]["n"])[0] if dist_riesgo else "N/D"
    puntaje_prom_global = df["Puntaje Total"].mean() if "Puntaje Total" in df.columns else 0

    parrafo = (
        f"En cumplimiento de la NOM-035-STPS-2018, se aplicó el cuestionario de la "
        f"Guía de Referencia II a un total de {total} trabajador{'es' if total!=1 else ''} "
        f"de la organización. El puntaje promedio global fue de {puntaje_prom_global:.1f} puntos "
        f"(escala 0-184), correspondiente al nivel de riesgo predominante '{nivel_predominante}' "
        f"conforme a la Tabla 5 de la norma. "
        + (f"Se identificaron {casos_criticos} caso(s) en nivel ALTO o MUY ALTO que "
           f"requieren intervención inmediata. " if casos_criticos > 0 else "")
        + (f"El dominio con mayor nivel de riesgo es '{dom_critico_nom}' "
           f"(nivel {dom_critico_niv}), " if dom_critico_niv not in ["NULO","BAJO"] else "")
        + (f"Se detectaron {alertas_violencia} caso(s) con indicadores de violencia laboral "
           f"que requieren atención prioritaria conforme al numeral 8.4.c de la norma. "
           if alertas_violencia > 0 else "")
        + (f"El área con mayor concentración de riesgo es '{foco_rojo}'. "
           if foco_rojo != "N/D" and seg_area.get(foco_rojo, {}).get("pct", 0) > 30 else "")
        + "El empleador deberá establecer las medidas de control correspondientes "
        f"y documentar las acciones realizadas conforme al numeral 8.4 de la NOM-035-STPS-2018."
    )

    return {
        "total":             total,
        "dist_riesgo":       dist_riesgo,
        "casos_criticos":    casos_criticos,
        "alertas_violencia": alertas_violencia,
        "promedios_dom":     promedios_dom,
        "niveles_dom":       niveles_dom,
        "dom_critico_id":    dom_critico_id,
        "dom_critico_nom":   dom_critico_nom,
        "dom_critico_niv":   dom_critico_niv,
        "viol_prom":         viol_prom,
        "seg_area":          dict(sorted(seg_area.items(),
                                         key=lambda x: x[1]["pct"], reverse=True)),
        "foco_rojo":         foco_rojo,
        "correlaciones":     correlaciones,
        "puntaje_prom":      round(puntaje_prom_global, 1),
        "nivel_predominante":nivel_predominante,
        "parrafo":           parrafo,
    }

# ══════════════════════════════════════════════════════════════════════════════
# EXCEL MEJORADO GUÍA II
# ══════════════════════════════════════════════════════════════════════════════
def _resolve_logo(logo_path: str) -> str:
    """Resolves logo path — tries absolute, then relative to this file."""
    if logo_path and os.path.exists(logo_path):
        return logo_path
    base = os.path.dirname(os.path.abspath(__file__)) if "__file__" in dir() else os.getcwd()
    alt  = os.path.join(base, logo_path) if logo_path else ""
    return alt if os.path.exists(alt) else ""

def generar_excel_g2(excel_path: str, cliente: str, razon: str) -> str:
    if not os.path.exists(excel_path):
        return None
    df = pd.read_excel(excel_path)
    if df.empty:
        return None

    fecha_str = datetime.now().strftime("%Y%m%d_%H%M")
    out = excel_path.replace(".xlsx", f"_reporte_g2_{fecha_str}.xlsx")

    with pd.ExcelWriter(out, engine="openpyxl") as w:
        df.to_excel(w, sheet_name="Datos", index=False)

    wb = load_workbook(out)

    # ── Hoja Datos — formato ──────────────────────────────────────────────────
    ws = wb["Datos"]
    for cell in ws[1]:
        cell.fill      = _fill(VERDE)
        cell.font      = _font(bold=True, size=10, color=BLANCO)
        cell.alignment = _align("center","center")
        cell.border    = _border(VERDE_MID)
    # Columna Nivel Riesgo — color por nivel
    nivel_col = None
    for idx, cell in enumerate(ws[1], 1):
        if cell.value == "Nivel Riesgo":
            nivel_col = idx
    viol_col = None
    for idx, cell in enumerate(ws[1], 1):
        if cell.value == "Alerta Violencia":
            viol_col = idx

    for ri, row in enumerate(ws.iter_rows(min_row=2), start=2):
        bg = "F0F6F0" if ri % 2 == 0 else BLANCO
        for cell in row:
            cell.fill = _fill(bg); cell.font = _font(size=9)
            cell.alignment = _align("left","center",wrap=True)
            cell.border = _border()
        if nivel_col:
            c = ws.cell(row=ri, column=nivel_col)
            col_txt, col_bg = NIVEL_COLORES.get(str(c.value or ""), ("#888","#eee"))
            c.fill = _fill(col_bg.replace("#",""))
            c.font = _font(bold=True, size=9, color=col_txt.replace("#",""))
        if viol_col:
            c = ws.cell(row=ri, column=viol_col)
            if "SÍ" in str(c.value or ""):
                c.fill = _fill(ROJO_CLARO)
                c.font = _font(bold=True, size=9, color=ROJO)

    _auto_width(ws)
    ws.freeze_panes = "A2"

    # ── Hoja Análisis ─────────────────────────────────────────────────────────
    a = _calcular_analiticas_g2(df)
    if not a:
        wb.save(out); return out

    ws_an = wb.create_sheet("Análisis")

    def _tit(fila, texto, color_bg):
        ws_an.merge_cells(f"A{fila}:L{fila}")
        ws_an[f"A{fila}"] = texto
        ws_an[f"A{fila}"].font      = _font(bold=True, size=11, color=BLANCO)
        ws_an[f"A{fila}"].fill      = _fill(color_bg)
        ws_an[f"A{fila}"].alignment = _align("left","center")
        ws_an.row_dimensions[fila].height = 22

    def _kv(fila, label, val, color_val=VERDE):
        ws_an[f"A{fila}"] = label
        ws_an[f"A{fila}"].font      = _font(size=10); ws_an[f"A{fila}"].fill = _fill(GRIS)
        ws_an[f"A{fila}"].border    = _border(); ws_an[f"A{fila}"].alignment = _align("left","center")
        ws_an[f"B{fila}"] = val
        ws_an[f"B{fila}"].font      = _font(bold=True, size=10, color=color_val)
        ws_an[f"B{fila}"].border    = _border(); ws_an[f"B{fila}"].alignment = _align("center","center")

    def _img(buf, celda, w=380, h=280):
        if buf:
            buf.seek(0)
            img = XLImage(buf); img.width = w; img.height = h
            ws_an.add_image(img, celda)

    # Título
    ws_an.merge_cells("A1:L1")
    ws_an["A1"] = "ANÁLISIS AVANZADO GUÍA II — NOM-035-STPS-2018"
    ws_an["A1"].font = _font(bold=True, size=14, color=BLANCO)
    ws_an["A1"].fill = _fill(VERDE)
    ws_an["A1"].alignment = _align("center","center")
    ws_an.row_dimensions[1].height = 30
    ws_an.merge_cells("A2:L2")
    ws_an["A2"] = f"Empresa: {razon}   |   Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ws_an["A2"].font = _font(size=9, color="888888")
    ws_an.row_dimensions[2].height = 16

    # BLOQUE 1: Distribución de riesgo — filas 4-12
    _tit(4, "1. DISTRIBUCIÓN POR NIVEL DE RIESGO (TABLA 5 NOM-035)", ROJO)
    hdrs = ["Nivel","Trabajadores","% del total","Acción requerida"]
    for ci, h in enumerate(hdrs, 1):
        c = ws_an.cell(row=5, column=ci, value=h)
        c.font = _font(bold=True, size=9, color=BLANCO)
        c.fill = _fill(VERDE); c.border = _border(VERDE_MID)
        c.alignment = _align("center","center")
    acciones = {
        "NULO":     "Sin acción requerida",
        "BAJO":     "Medidas preventivas",
        "MEDIO":    "Programa de intervención (obligatorio)",
        "ALTO":     "Intervención inmediata + seguimiento",
        "MUY ALTO": "URGENTE — Intervención inmediata",
    }
    for ri, nv in enumerate(["NULO","BAJO","MEDIO","ALTO","MUY ALTO"], start=6):
        d = a["dist_riesgo"].get(nv, {"n":0,"pct":0})
        bg = NIVEL_COLORES.get(nv, ("#888","#eee"))[1].replace("#","")
        for ci, val in enumerate([nv, d["n"], f"{d['pct']}%", acciones[nv]], 1):
            c = ws_an.cell(row=ri, column=ci, value=val)
            c.fill = _fill(bg); c.border = _border()
            c.font = _font(size=9, bold=(nv in ["ALTO","MUY ALTO"]))
            c.alignment = _align("center","center")
    _img(_graf_distribucion_riesgo(df), "F4", 360, 280)

    # BLOQUE 2: Mapa de calor dominios — filas 14-26
    _tit(14, "2. MAPA DE CALOR POR DOMINIOS", "C8A600")
    hdrs_d = ["Dominio","Prom. pts","Nivel","% del máximo"]
    for ci, h in enumerate(hdrs_d, 1):
        c = ws_an.cell(row=15, column=ci, value=h)
        c.font = _font(bold=True, size=9, color=BLANCO)
        c.fill = _fill(VERDE); c.border = _border()
        c.alignment = _align("center","center")
    maximos = {0:16,1:48,2:24,3:12,4:20,5:36,6:16,7:12}
    for ri, dom_id in enumerate(range(8), start=16):
        prom = a["promedios_dom"].get(dom_id, 0)
        nv   = a["niveles_dom"].get(dom_id, "NULO")
        pct  = round(prom/maximos[dom_id]*100, 1) if maximos[dom_id] > 0 else 0
        bg   = NIVEL_COLORES.get(nv, ("#888","#eee"))[1].replace("#","")
        for ci, val in enumerate([NOMBRES_DOM[dom_id], prom, nv, f"{pct}%"], 1):
            c = ws_an.cell(row=ri, column=ci, value=val)
            c.fill = _fill(bg); c.border = _border()
            c.font = _font(size=9, bold=(nv in ["ALTO","MUY ALTO"]))
            c.alignment = _align("center","center")
    _img(_graf_mapa_calor_dominios(df), "F14", 440, 300)

    # BLOQUE 3: Violencia laboral — filas 28-36
    _tit(28, "3. ANÁLISIS DE VIOLENCIA LABORAL (ítems 44-46) — FOCO ROJO", ROJO)
    ws_an.merge_cells("A29:L29")
    c_viol = ws_an["A29"]
    c_viol.value = (f"Trabajadores con alerta de violencia: {a['alertas_violencia']} "
                    f"de {a['total']} ({round(a['alertas_violencia']/a['total']*100,1) if a['total']>0 else 0}%)")
    c_viol.font = _font(bold=True, size=10, color=ROJO if a["alertas_violencia"] > 0 else VERDE)
    c_viol.fill = _fill(ROJO_CLARO if a["alertas_violencia"] > 0 else "D6E4D8")
    c_viol.border = _border()
    items_viol = [
        ("P34 — Críticas constantes",              a["viol_prom"].get("P34", 0)),
        ("P35 — Burlas, humillaciones",            a["viol_prom"].get("P35", 0)),
        ("P36 — Me excluyen",                      a["viol_prom"].get("P36", 0)),
        ("P37 — Me hacen parecer mal trabajador",  a["viol_prom"].get("P37", 0)),
        ("P38 — Ignoran mis éxitos",               a["viol_prom"].get("P38", 0)),
        ("P39 — Bloquean ascensos",                a["viol_prom"].get("P39", 0)),
        ("P40 — Presencié violencia",              a["viol_prom"].get("P40", 0)),
    ]
    for ri, (etiq, prom) in enumerate(items_viol, start=30):
        for ci, val in enumerate([etiq, f"{prom:.2f}/4.0",
                                   "URGENTE" if prom > 1 else "Monitorear" if prom > 0 else "OK"], 1):
            c = ws_an.cell(row=ri, column=ci, value=val)
            bg = ROJO_CLARO if prom > 1 else AMARILLO_BG if prom > 0 else "D6E4D8"
            c.fill = _fill(bg); c.border = _border()
            c.font = _font(size=9, bold=(prom > 1))
            c.alignment = _align("center","center")
    _img(_graf_violencia(df), "F28", 380, 240)

    # BLOQUE 4: Segmentación por área — filas 38-50
    _tit(38, "4. SEGMENTACIÓN POR ÁREA — FOCOS DE ATENCIÓN", "69A2D8")
    if a["seg_area"]:
        hdrs_a = ["Área","Total","Críticos","% críticos","Ptj.Prom","Nivel prom","Prioridad"]
        for ci, h in enumerate(hdrs_a, 1):
            c = ws_an.cell(row=39, column=ci, value=h)
            c.font = _font(bold=True, size=9, color=BLANCO)
            c.fill = _fill(VERDE); c.border = _border()
            c.alignment = _align("center","center")
        for ri, (area, dat) in enumerate(a["seg_area"].items(), start=40):
            prio = "🔴 INMEDIATA" if dat["pct"]>=50 else "🟡 ALTA" if dat["pct"]>=25 else "🟢 SEGUIMIENTO"
            bg_a = ROJO_CLARO if dat["pct"]>=50 else AMARILLO_BG if dat["pct"]>=25 else "D6E4D8"
            for ci, val in enumerate([area,dat["total"],dat["criticos"],
                                       f"{dat['pct']}%",dat["puntaje_prom"],
                                       dat["nivel_prom"],prio], 1):
                c = ws_an.cell(row=ri, column=ci, value=val)
                c.fill = _fill(bg_a); c.border = _border()
                c.font = _font(size=9); c.alignment = _align("center","center")
        _img(_graf_por_area(df), "H38", 400, 280)

    # BLOQUE 5: Párrafo ejecutivo — últimas filas
    fila_final = 54
    _tit(fila_final, "5. PÁRRAFO EJECUTIVO DE CUMPLIMIENTO LEGAL (Num. 10.2 NOM-035)", VERDE)
    ws_an.merge_cells(f"A{fila_final+1}:L{fila_final+6}")
    c_par = ws_an[f"A{fila_final+1}"]
    c_par.value     = a["parrafo"]
    c_par.font      = _font(size=10)
    c_par.fill      = _fill("F0F6F0")
    c_par.alignment = _align("left","top",wrap=True)
    c_par.border    = _border(VERDE)
    ws_an.row_dimensions[fila_final+1].height = 90

    # Anchos
    for col, w in [("A",32),("B",14),("C",14),("D",12),("E",12),
                   ("F",12),("G",12),("H",12),("I",12),("J",12),("K",12),("L",12)]:
        ws_an.column_dimensions[col].width = w

    # ── Hoja Gráficas ─────────────────────────────────────────────────────────
    ws_g = wb.create_sheet("Gráficas")
    ws_g.merge_cells("A1:P1")
    ws_g["A1"] = "GRÁFICAS DE ANÁLISIS — GUÍA II NOM-035-STPS-2018"
    ws_g["A1"].font = _font(bold=True, size=14, color=BLANCO)
    ws_g["A1"].fill = _fill(VERDE); ws_g["A1"].alignment = _align("center","center")
    ws_g.row_dimensions[1].height = 30

    for c in range(1, 17):
        ws_g.column_dimensions[get_column_letter(c)].width = 12

    graf_layout = [
        (_graf_distribucion_riesgo(df), "B3",  "DISTRIBUCIÓN POR NIVEL DE RIESGO", "B2"),
        (_graf_mapa_calor_dominios(df),  "J3",  "MAPA DE CALOR POR DOMINIOS",       "J2"),
        (_graf_radar_dominios(df),       "B23", "RADAR DE DOMINIOS",                "B22"),
        (_graf_violencia(df),            "J23", "ANÁLISIS DE VIOLENCIA LABORAL",     "J22"),
        (_graf_por_area(df),             "B43", "RIESGO PROMEDIO POR ÁREA",          "B42"),
        (_graf_carga_trabajo(df),        "J43", "ANÁLISIS CARGA DE TRABAJO",         "J42"),
    ]
    for buf, celda_img, titulo, celda_tit in graf_layout:
        ws_g[celda_tit] = titulo
        ws_g[celda_tit].font = _font(bold=True, size=9, color=BLANCO)
        ws_g[celda_tit].fill = _fill(VERDE_MID)
        ws_g[celda_tit].alignment = _align("center","center")
        col_l  = celda_tit[0]
        row_t  = int(celda_tit[1:])
        col_fi = get_column_letter(ord(col_l) - ord('A') + 1 + 7)
        try: ws_g.merge_cells(f"{celda_tit}:{col_fi}{row_t}")
        except: pass
        ws_g.row_dimensions[row_t].height = 22
        if buf:
            buf.seek(0); img = XLImage(buf); img.width=420; img.height=320
            ws_g.add_image(img, celda_img)

    wb.save(out)
    return out

# ══════════════════════════════════════════════════════════════════════════════
# WORD GUÍA II
# ══════════════════════════════════════════════════════════════════════════════
def generar_word_g2(excel_path: str, cliente: str, razon: str,
                    logo_rf: str = None, logo_cliente: str = None) -> str:
    logo_rf      = _resolve_logo(logo_rf)      if logo_rf      else ""
    logo_cliente = _resolve_logo(logo_cliente) if logo_cliente else ""
    if not os.path.exists(excel_path):
        return None
    df = pd.read_excel(excel_path)
    if df.empty:
        return None

    fecha_str = datetime.now().strftime("%Y%m%d_%H%M")
    meses_es  = {"January":"enero","February":"febrero","March":"marzo",
                 "April":"abril","May":"mayo","June":"junio","July":"julio",
                 "August":"agosto","September":"septiembre","October":"octubre",
                 "November":"noviembre","December":"diciembre"}
    fd = datetime.now()
    fecha_es = f"{fd.day} de {meses_es.get(fd.strftime('%B'), fd.strftime('%B'))} de {fd.year}"
    out = excel_path.replace(".xlsx", f"_informe_g2_{fecha_str}.docx")

    doc    = Document()
    _fig_n = [0]

    for sec in doc.sections:
        sec.top_margin    = Cm(2.5); sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5); sec.right_margin  = Cm(2.5)
        sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)

    nml = doc.styles["Normal"]
    nml.font.name = "Calibri"; nml.font.size = Pt(11)
    nml.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _run(p, text, bold=False, size=11, color_hex=None, italic=False):
        r = p.add_run(text)
        r.bold=bold; r.italic=italic; r.font.name="Calibri"; r.font.size=Pt(size)
        if color_hex: r.font.color.rgb = RGBColor.from_string(color_hex)
        return r

    def _p(text="", bold=False, size=11, color_hex=None,
           align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6, italic=False):
        p = doc.add_paragraph()
        p.alignment=align; p.paragraph_format.space_before=Pt(sb)
        p.paragraph_format.space_after=Pt(sa)
        if text: _run(p, text, bold=bold, size=size, color_hex=color_hex, italic=italic)
        return p

    def _sec_tit(texto, color=VERDE):
        p = doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"8")
        bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),color)
        pBdr.append(bot); pPr.append(pBdr)
        _run(p, texto, bold=True, size=13, color_hex=color)
        return p

    def _subsec(texto, color=VERDE_MID):
        p = doc.add_paragraph()
        p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
        _run(p, texto, bold=True, size=11, color_hex=color)
        return p

    def _figura(buf, caption, width_cm=13):
        if not buf: return
        buf.seek(0); _fig_n[0] += 1
        pi = doc.add_paragraph()
        pi.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pi.paragraph_format.space_before=Pt(6); pi.paragraph_format.space_after=Pt(2)
        pi.add_run().add_picture(buf, width=Cm(width_cm))
        pc = doc.add_paragraph()
        pc.alignment=WD_ALIGN_PARAGRAPH.CENTER; pc.paragraph_format.space_after=Pt(8)
        _run(pc, f"Figura {_fig_n[0]}. {caption}", size=9, italic=True, color_hex="606060")

    def _set_bg(cell, hex_c):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
        shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_c); tcPr.append(shd)

    def _set_bdr(cell, color=VERDE):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        tcB=OxmlElement("w:tcBorders")
        for side in ["top","left","bottom","right"]:
            el=OxmlElement(f"w:{side}"); el.set(qn("w:val"),"single")
            el.set(qn("w:sz"),"4"); el.set(qn("w:space"),"0")
            el.set(qn("w:color"),color); tcB.append(el)
        tcPr.append(tcB)

    def _valign(cell, val="center"):
        tcPr=cell._tc.get_or_add_tcPr()
        vAl=OxmlElement("w:vAlign"); vAl.set(qn("w:val"),val); tcPr.append(vAl)

    def _tbl_cell(cell, text, bold=False, size=10, color_hex=None,
                  align=WD_ALIGN_PARAGRAPH.CENTER, bg=None, bdr=VERDE):
        if bg: _set_bg(cell, bg)
        _set_bdr(cell, bdr); _valign(cell)
        p=cell.paragraphs[0]; p.alignment=align
        p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
        _run(p, text, bold=bold, size=size, color_hex=color_hex)

    # ── Encabezado NOM-004 ────────────────────────────────────────────────────
    def _build_header(section):
        section.different_first_page_header_footer = True
        fph = section.first_page_header
        if not fph.paragraphs: fph.add_paragraph()
        fph.paragraphs[0].text = ""

        for footer_obj in [section.footer, section.first_page_footer]:
            for p in footer_obj.paragraphs:
                for r in p.runs: r.text = ""
            fp = footer_obj.paragraphs[0] if footer_obj.paragraphs else footer_obj.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fp.paragraph_format.space_before=Pt(0); fp.paragraph_format.space_after=Pt(0)
            pPr=fp._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
            top_el=OxmlElement("w:top"); top_el.set(qn("w:val"),"single")
            top_el.set(qn("w:sz"),"4"); top_el.set(qn("w:space"),"1")
            top_el.set(qn("w:color"),"D0D0D0"); pBdr.append(top_el); pPr.append(pBdr)
            r_pie = fp.add_run("DOCUMENTO ORIGINAL")
            r_pie.font.name="Calibri"; r_pie.font.size=Pt(8)
            r_pie.font.color.rgb=RGBColor(96,96,96)

        hdr = section.header
        for p in hdr.paragraphs:
            for r in p.runs: r.text = ""
        tbl = hdr.add_table(rows=3, cols=3, width=Cm(16))
        tbl.style = "Table Grid"
        COL_W = [3, 10, 3]
        for ri in range(3):
            for ci in range(3):
                tcPr=tbl.cell(ri,ci)._tc.get_or_add_tcPr()
                tcW=OxmlElement("w:tcW"); tcW.set(qn("w:w"),str(int(COL_W[ci]*567)))
                tcW.set(qn("w:type"),"dxa"); tcPr.append(tcW)
        tbl.cell(0,0).merge(tbl.cell(2,0))
        tbl.cell(0,2).merge(tbl.cell(1,2))

        c_logo = tbl.cell(0,0)
        _set_bg(c_logo, BLANCO); _set_bdr(c_logo, VERDE); _valign(c_logo,"center")
        p_l=c_logo.paragraphs[0]; p_l.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p_l.paragraph_format.space_before=Pt(2); p_l.paragraph_format.space_after=Pt(2)
        if logo_cliente and os.path.exists(logo_cliente):
            try: p_l.add_run().add_picture(logo_cliente, width=Cm(2.6))
            except: _run(p_l, cliente, bold=True, size=9, color_hex=VERDE)

        c_emp = tbl.cell(0,1)
        _set_bg(c_emp, BLANCO); _set_bdr(c_emp, VERDE); _valign(c_emp,"center")
        p_emp=c_emp.paragraphs[0]; p_emp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p_emp.paragraph_format.space_before=Pt(3); p_emp.paragraph_format.space_after=Pt(3)
        _run(p_emp, "Empresa:  ", size=7.5, color_hex="888888")
        _run(p_emp, razon, size=8, bold=True, color_hex=VERDE)

        c_tit = tbl.cell(1,1)
        _set_bg(c_tit, VERDE_CLARO); _set_bdr(c_tit, VERDE); _valign(c_tit,"center")
        p_t=c_tit.paragraphs[0]; p_t.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p_t.paragraph_format.space_before=Pt(3); p_t.paragraph_format.space_after=Pt(3)
        _run(p_t, "NOM-035-STPS-2018 — GUÍA II  |  FACTORES DE RIESGO PSICOSOCIAL "
             "Y EVALUACIÓN DEL ENTORNO ORGANIZACIONAL",
             size=7.5, bold=True, color_hex=VERDE)

        c_fec = tbl.cell(2,1)
        _set_bg(c_fec, BLANCO); _set_bdr(c_fec, VERDE); _valign(c_fec,"center")
        p_f=c_fec.paragraphs[0]; p_f.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p_f.paragraph_format.space_before=Pt(3); p_f.paragraph_format.space_after=Pt(3)
        _run(p_f, "Fecha de emisión:  ", size=7.5, color_hex="888888")
        _run(p_f, fecha_es, size=7.5, bold=True)

        c_pag = tbl.cell(2,2)
        _set_bg(c_pag, BLANCO); _set_bdr(c_pag, VERDE); _valign(c_pag,"center")
        p_pag=c_pag.paragraphs[0]; p_pag.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p_pag.paragraph_format.space_before=Pt(3); p_pag.paragraph_format.space_after=Pt(3)
        _run(p_pag, "No. Página:  ", size=7.5, color_hex="888888")
        fld1=OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"),"begin")
        ins=OxmlElement("w:instrText"); ins.text=" PAGE "
        fld2=OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"),"end")
        r_pg=p_pag.add_run()
        r_pg._r.append(fld1); r_pg._r.append(ins); r_pg._r.append(fld2)
        r_pg.font.name="Calibri"; r_pg.font.size=Pt(8); r_pg.bold=True

    _build_header(doc.sections[0])

    # ── PORTADA ───────────────────────────────────────────────────────────────
    tbl_port = doc.add_table(rows=1, cols=3)
    tbl_port.alignment = WD_TABLE_ALIGNMENT.CENTER
    if logo_rf and os.path.exists(logo_rf):
        c=tbl_port.cell(0,0); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
        c.paragraphs[0].add_run().add_picture(logo_rf, width=Cm(3.5))
    sep_r=tbl_port.cell(0,1).paragraphs[0].add_run("|")
    sep_r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC); sep_r.font.size=Pt(28)
    tbl_port.cell(0,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if logo_cliente and os.path.exists(logo_cliente):
        c=tbl_port.cell(0,2); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
        c.paragraphs[0].add_run().add_picture(logo_cliente, width=Cm(3.5))
    _p(sa=4)

    p_ln=doc.add_paragraph(); p_ln.paragraph_format.space_after=Pt(0)
    pPr=p_ln._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    bot=OxmlElement("w:bottom"); bot.set(qn("w:val"),"single")
    bot.set(qn("w:sz"),"12"); bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),VERDE)
    pBdr.append(bot); pPr.append(pBdr)

    _p("NOM-035-STPS-2018", bold=True, size=22, color_hex=VERDE,
       align=WD_ALIGN_PARAGRAPH.CENTER, sb=14, sa=3)
    _p("GUÍA DE REFERENCIA II", bold=True, size=13, color_hex=VERDE_MID,
       align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=5)
    _p("Identificación y Análisis de los Factores de Riesgo Psicosocial "
       "y Evaluación del Entorno Organizacional",
       size=11, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4)
    _p("(Centros de trabajo con 16 a 50 trabajadores · 46 preguntas)",
       size=10, color_hex="888888", align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=14)

    a = _calcular_analiticas_g2(df)
    total = a.get("total", len(df))

    tbl_dat = doc.add_table(rows=3, cols=2)
    tbl_dat.style = "Table Grid"; tbl_dat.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(lbl,val) in enumerate([("Empresa:", razon),
                                   ("Fecha:", fecha_es),
                                   ("Elaboró:", "RFRANYUTTI, CONCIENCIA VERDE Y LABORAL S.C.")]):
        c0=tbl_dat.cell(i,0); c1=tbl_dat.cell(i,1)
        _set_bg(c0,VERDE_CLARO); _set_bdr(c0,VERDE)
        _set_bg(c1,BLANCO);      _set_bdr(c1,VERDE)
        _tbl_cell(c0, lbl, bold=True, size=11, color_hex=VERDE, align=WD_ALIGN_PARAGRAPH.LEFT, bg=None)
        _tbl_cell(c1, val, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bg=None)
    _p(sa=8)

    # Estadísticas portada
    casos_c = a.get("casos_criticos", 0)
    alrts_v = a.get("alertas_violencia", 0)
    tbl_st  = doc.add_table(rows=1, cols=3)
    tbl_st.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(num,etiq,col) in enumerate([
        (str(total), "Trabajadores evaluados", VERDE),
        (str(casos_c), "Casos Alto/Muy Alto", ROJO if casos_c>0 else VERDE),
        (str(alrts_v), "Alertas de violencia", ROJO if alrts_v>0 else VERDE),
    ]):
        c=tbl_st.cell(0,i); _set_bg(c,col); _set_bdr(c,col); _valign(c,"center")
        pn=c.add_paragraph(); pn.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(pn, num, bold=True, size=28, color_hex=BLANCO)
        pe=c.add_paragraph(); pe.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pe.paragraph_format.space_after=Pt(4)
        _run(pe, etiq, size=9, color_hex=BLANCO)

    doc.add_page_break()

    # ── SECCIÓN 1 — RESULTADOS GRÁFICOS ──────────────────────────────────────
    _sec_tit("1. RESULTADOS GRÁFICOS")
    _p("Las siguientes gráficas presentan el análisis de los resultados conforme "
       "a los criterios de la Guía II de la NOM-035-STPS-2018.", sa=8)
    _figura(_graf_distribucion_riesgo(df), "Distribución por Nivel de Riesgo (Tabla 5 NOM-035)", 12)
    _figura(_graf_radar_dominios(df),       "Radar de Dominios — Porcentaje del Puntaje Máximo", 10)
    _figura(_graf_mapa_calor_dominios(df),  "Mapa de Calor por Dominios",                        13)

    doc.add_page_break()

    # ── SECCIÓN 2 — TABLA RESULTADOS ─────────────────────────────────────────
    _sec_tit("2. TABLA DE RESULTADOS POR TRABAJADOR")
    _p(f"Resultados individuales de los {total} trabajadores evaluados.", sa=6)
    COLS_W = ["Folio","Nombre","Puesto","Área","Puntaje Total","Nivel Riesgo","Alerta Violencia"]
    cols_d = [c for c in COLS_W if c in df.columns]
    tbl_r  = doc.add_table(rows=1, cols=len(cols_d))
    tbl_r.style = "Table Grid"; tbl_r.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,col in enumerate(cols_d):
        _tbl_cell(tbl_r.rows[0].cells[i], col, bold=True, size=9,
                  color_hex=BLANCO, bg=VERDE)
    for ri,(_,row) in enumerate(df.iterrows()):
        cells=tbl_r.add_row().cells; bg_r="F4F4F4" if ri%2==0 else BLANCO
        for i,col in enumerate(cols_d):
            val=str(row.get(col,"")) if pd.notna(row.get(col,"")) else ""
            c=cells[i]
            if col=="Nivel Riesgo":
                col_t, col_bg = NIVEL_COLORES.get(val, ("#888","#eee"))
                _tbl_cell(c, val, bold=(val in ["ALTO","MUY ALTO"]), size=9,
                          color_hex=col_t.replace("#",""), bg=col_bg.replace("#",""))
            elif col=="Alerta Violencia" and "SÍ" in val:
                _tbl_cell(c, val, bold=True, size=9, color_hex=ROJO, bg=ROJO_CLARO)
            else:
                _tbl_cell(c, val, size=9, bg=bg_r)

    doc.add_page_break()

    # ── SECCIÓN 3 — ANÁLISIS POR DOMINIOS ────────────────────────────────────
    _sec_tit("3. ANÁLISIS POR DOMINIOS")
    _p("El análisis por dominios permite identificar las áreas específicas donde "
       "se concentran los factores de riesgo psicosocial.", sa=6)

    maximos = {0:16,1:48,2:24,3:12,4:20,5:36,6:16,7:12}
    tbl_dom = doc.add_table(rows=1, cols=4)
    tbl_dom.style = "Table Grid"; tbl_dom.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(["Dominio","Puntaje Promedio","Nivel","% del máximo"]):
        _tbl_cell(tbl_dom.rows[0].cells[i], h, bold=True, size=10,
                  color_hex=BLANCO, bg=VERDE)
    for dom_id in range(8):
        prom = a["promedios_dom"].get(dom_id, 0)
        nv   = a["niveles_dom"].get(dom_id, "NULO")
        pct  = round(prom/maximos[dom_id]*100,1)
        col_t, col_bg = NIVEL_COLORES.get(nv, ("#888","#eee"))
        cells=tbl_dom.add_row().cells
        _tbl_cell(cells[0], NOMBRES_DOM[dom_id], size=10, bg=col_bg.replace("#",""))
        _tbl_cell(cells[1], str(prom), size=10, bold=True,
                  color_hex=col_t.replace("#",""), bg=col_bg.replace("#",""))
        _tbl_cell(cells[2], nv, size=10, bold=(nv in ["ALTO","MUY ALTO"]),
                  color_hex=col_t.replace("#",""), bg=col_bg.replace("#",""))
        _tbl_cell(cells[3], f"{pct}%", size=10, bg=col_bg.replace("#",""))

    _p(sa=6)
    _p(f"El dominio con mayor nivel de riesgo es "
       f"'{a['dom_critico_nom']}' (nivel {a['dom_critico_niv']}).",
       bold=True, color_hex=ROJO if a["dom_critico_niv"] in ["ALTO","MUY ALTO"] else VERDE, sa=8)

    # Análisis específicos
    _subsec("3.1 Carga de Trabajo (ítems 5-16)", AMBAR)
    _p("Los ítems 5-7 miden cantidad de trabajo, 8-10 velocidad y carga mental, "
       "11-14 responsabilidades y reacciones emocionales. "
       "Los ítems 15-16 evalúan el desarrollo de habilidades (escala inversa).", sa=4)
    _figura(_graf_carga_trabajo(df), "Análisis Detallado — Carga de Trabajo (ítems 5-16)", 13)

    _subsec("3.2 Falta de Control sobre el Trabajo (ítems 17-22)", AZUL)
    _p("Estos ítems evalúan si el trabajador recibe instrucciones claras (17-19), "
       "puede organizar su trabajo con autonomía (20-21) y recibe capacitación suficiente (22). "
       "Puntajes altos indican falta de control y autonomía.", sa=6)

    _subsec("3.3 Interferencia Trabajo-Familia (ítems 26-30)", AMBAR)
    _p("Cuando el trabajador reporta que el trabajo invade su tiempo familiar "
       "o que sigue pensando en el trabajo en casa (ítems 28-30), "
       "se recomienda implementar políticas de desconexión digital "
       "y horarios flexibles.", sa=6)

    _subsec("3.4 Liderazgo (ítems 31-39)", VERDE_MID)
    _p("El liderazgo positivo (jefe que escucha, comunica, distribuye trabajo equitativamente) "
       "actúa como factor protector. Puntajes bajos en estos ítems indican problemas "
       "de supervisión que pueden amplificar otros factores de riesgo.", sa=6)

    doc.add_page_break()

    # ── SECCIÓN 4 — VIOLENCIA LABORAL ────────────────────────────────────────
    _sec_tit("4. ANÁLISIS DE VIOLENCIA LABORAL (ítems 44-46)", ROJO)

    if a["alertas_violencia"] > 0:
        tbl_alerta = doc.add_table(rows=1, cols=1)
        tbl_alerta.alignment = WD_TABLE_ALIGNMENT.CENTER
        c_al = tbl_alerta.cell(0,0)
        _set_bg(c_al, ROJO_CLARO); _set_bdr(c_al, ROJO)
        p_al = c_al.paragraphs[0]; p_al.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_al.paragraph_format.space_before=Pt(6); p_al.paragraph_format.space_after=Pt(6)
        _run(p_al, f"🚨 ALERTA — {a['alertas_violencia']} CASO(S) DE VIOLENCIA LABORAL DETECTADO(S)",
             bold=True, size=11, color_hex=ROJO)
        _p(sa=6)

    _p("Los ítems 44 al 46 son los más delicados del cuestionario. "
       "Cualquier respuesta positiva (distinta de 'Nunca') debe activar "
       "el protocolo de atención conforme al numeral 8.4.c de la NOM-035-STPS-2018.", sa=4)

    for col, desc in [("P44","Me ignoran o excluyen"),
                       ("P45","Burlas, humillaciones, ridiculizaciones"),
                       ("P46","Responsabilidades sin autoridad / tareas sin recursos")]:
        prom = a["viol_prom"].get(col, 0)
        nivel_v = "URGENTE" if prom > 1 else "MONITOREAR" if prom > 0 else "SIN ALERTA"
        color_v = ROJO if prom > 1 else AMBAR if prom > 0 else VERDE
        pb = doc.add_paragraph(style="List Bullet")
        pb.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pb.paragraph_format.space_after = Pt(3)
        _run(pb, f"{col} — {desc}: ", bold=True, size=11)
        _run(pb, f"Promedio {prom:.2f}/4.0 — {nivel_v}", size=11, color_hex=color_v)

    _p(sa=4)
    _figura(_graf_violencia(df), "Análisis de Violencia Laboral — Promedio por Ítem", 11)

    _p(sa=6)
    _subsec("Protocolo de actuación inmediata (cuando hay alerta):", ROJO)
    for accion in [
        "Identificar al trabajador afectado y garantizar su confidencialidad.",
        "Documentar el caso e iniciar investigación interna conforme al Código de Ética.",
        "Aplicar medidas de control inmediatas: separar a las partes involucradas si aplica.",
        "Canalizar al trabajador a apoyo psicológico especializado.",
        "Revisar y actualizar el Código de Ética y políticas anti-acoso de la empresa.",
        "Dar seguimiento al caso y reportar acciones tomadas en el expediente NOM-035.",
    ]:
        pb = doc.add_paragraph(style="List Bullet")
        pb.paragraph_format.space_after=Pt(2)
        _run(pb, accion, size=11)

    doc.add_page_break()

    # ── SECCIÓN 5 — SEGMENTACIÓN POR ÁREA ───────────────────────────────────
    _sec_tit("5. SEGMENTACIÓN POR ÁREA — FOCOS DE ATENCIÓN")
    _p("El cruce de datos por área permite identificar si los factores de riesgo "
       "están concentrados en departamentos específicos, lo que permite priorizar "
       "las intervenciones.", sa=6)
    if a["seg_area"]:
        _p(f"El área con mayor concentración de riesgo es "
           f"'{a['foco_rojo']}'. "
           f"Semáforo: 🔴 ≥50% casos críticos — 🟡 25-49% — 🟢 <25%", sa=4)
        _figura(_graf_por_area(df), "Puntaje Promedio por Área de Trabajo", 13)

    if a["correlaciones"]:
        _subsec("Patrones y correlaciones detectadas:", AMBAR)
        for corr in a["correlaciones"]:
            pb = doc.add_paragraph(style="List Bullet")
            pb.paragraph_format.space_after=Pt(3)
            _run(pb, corr, size=11)

    doc.add_page_break()

    # ── SECCIÓN 6 — PROGRAMA DE ACCIÓN (Num. 8.4) ────────────────────────────
    _sec_tit("6. PROGRAMA DE ACCIÓN (Numeral 8.4 NOM-035-STPS-2018)")
    _p("Cuando los resultados indican nivel de riesgo medio o superior, la norma "
       "obliga al empleador a establecer un programa de intervención. "
       "A continuación se presentan las medidas recomendadas por nivel.", sa=6)

    dominios_criticos = [NOMBRES_DOM[d] for d, nv in a["niveles_dom"].items()
                         if nv in ["ALTO","MUY ALTO"]]
    dominios_medios   = [NOMBRES_DOM[d] for d, nv in a["niveles_dom"].items()
                         if nv == "MEDIO"]

    _subsec("Nivel 1 — Acciones organizacionales (políticas):", VERDE)
    nivel1 = [
        "Revisar y ajustar la distribución de cargas de trabajo para garantizar equidad.",
        "Implementar política de desconexión digital fuera del horario laboral.",
        "Establecer canales formales de comunicación jefe-colaborador con retroalimentación periódica.",
        "Actualizar el Reglamento Interior de Trabajo incluyendo cero tolerancia a violencia laboral.",
        "Garantizar que cada trabajador tenga por escrito sus funciones, responsabilidades y a quién reportar.",
    ]
    if dominios_criticos:
        nivel1.append(f"Atención prioritaria en los dominios críticos: {', '.join(dominios_criticos)}.")
    for a1 in nivel1:
        pb=doc.add_paragraph(style="List Bullet"); pb.paragraph_format.space_after=Pt(3)
        _run(pb, a1, size=11)

    _p(sa=4)
    _subsec("Nivel 2 — Acciones grupales (talleres y capacitación):", AMBAR)
    nivel2 = [
        "Taller de manejo del estrés y técnicas de relajación (si predomina carga de trabajo alta).",
        "Taller de comunicación asertiva y resolución de conflictos (si hay problemas de liderazgo).",
        "Capacitación en prevención de violencia laboral y acoso dirigida a supervisores y trabajadores.",
        "Programa de integración de equipos para fortalecer el compañerismo y las relaciones laborales.",
        "Sesiones grupales de sensibilización sobre el equilibrio vida-trabajo.",
    ]
    for a2 in nivel2:
        pb=doc.add_paragraph(style="List Bullet"); pb.paragraph_format.space_after=Pt(3)
        _run(pb, a2, size=11)

    _p(sa=4)
    _subsec("Nivel 3 — Atención individual (si aplica):", ROJO)
    nivel3 = [
        "Canalización a psicólogo o médico para trabajadores en nivel MUY ALTO.",
        "Seguimiento personalizado a los casos con alertas de violencia laboral.",
        "Orientación individual sobre derechos laborales y canales de denuncia.",
        "Evaluación médica especializada para trabajadores con síntomas de estrés severo.",
    ]
    for a3 in nivel3:
        pb=doc.add_paragraph(style="List Bullet"); pb.paragraph_format.space_after=Pt(3)
        _run(pb, a3, size=11)

    doc.add_page_break()

    # ── SECCIÓN 7 — INFORME PARA INSPECCIÓN STPS (Num. 10.2) ─────────────────
    _sec_tit("7. INFORME DE RESULTADOS (Numeral 10.2 NOM-035-STPS-2018)")
    _p("Este informe se genera conforme a los requisitos del numeral 10.2 de la "
       "NOM-035-STPS-2018 para su presentación ante la STPS en caso de inspección.", sa=8)

    secciones_informe = [
        ("7.1 Datos del centro de trabajo", [
            f"Empresa: {razon}",
            f"Fecha de evaluación: {fecha_es}",
            "Guía aplicada: Guía de Referencia II (16-50 trabajadores)",
            f"Total de trabajadores evaluados: {total}",
            "Responsable de la evaluación: RFRANYUTTI, CONCIENCIA VERDE Y LABORAL S.C.",
        ]),
        ("7.2 Objetivo", [
            "Identificar los factores de riesgo psicosocial presentes en el centro de trabajo.",
            "Evaluar el entorno organizacional mediante la aplicación del cuestionario oficial.",
            "Determinar el nivel de riesgo conforme a la Tabla 5 de la NOM-035-STPS-2018.",
        ]),
        ("7.3 Método aplicado", [
            "Instrumento: Cuestionario Guía II de la NOM-035-STPS-2018 (46 ítems).",
            "Escala de respuesta: Siempre (4), Casi siempre (3), Algunas veces (2), "
            "Casi nunca (1), Nunca (0).",
            "Aplicación: Individual, digital, con garantía de confidencialidad.",
            "Procesamiento: Cálculo de puntaje total y por dominio conforme a la norma.",
        ]),
        ("7.4 Resultados por categoría", [
            f"Nivel de riesgo predominante: {a.get('nivel_predominante','N/D')}",
            f"Puntaje promedio global: {a.get('puntaje_prom', 0)} puntos",
            f"Trabajadores en nivel ALTO o MUY ALTO: {a.get('casos_criticos',0)}",
            f"Dominio con mayor riesgo: {a.get('dom_critico_nom','N/D')} (nivel {a.get('dom_critico_niv','N/D')})",
            f"Alertas de violencia laboral: {a.get('alertas_violencia',0)}",
        ]),
    ]

    for subtit, puntos in secciones_informe:
        _subsec(subtit, VERDE)
        for punto in puntos:
            pb=doc.add_paragraph(style="List Bullet"); pb.paragraph_format.space_after=Pt(2)
            _run(pb, punto, size=11)
        _p(sa=4)

    # Conclusiones
    _subsec("7.5 Conclusiones", VERDE)
    _p(a.get("parrafo",""), sa=8)

    # Firma
    _p(sa=20)
    tbl_firma = doc.add_table(rows=2, cols=2)
    tbl_firma.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_emp_firma = tbl_firma.cell(0,0)
    c_emp_firma.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(c_emp_firma.paragraphs[0], "_"*40, size=11)
    p2 = c_emp_firma.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, "Representante del Empleador", size=9, color_hex="888888")

    c_resp = tbl_firma.cell(0,1)
    c_resp.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(c_resp.paragraphs[0], "_"*40, size=11)
    p3 = c_resp.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p3, "RFRANYUTTI — Responsable de la Evaluación", size=9, color_hex="888888")

    doc.save(out)
    return out
